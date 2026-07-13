from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(r"C:\claudcode_ap\growth_lab_core")
OUT_DIR = ROOT / "docs_export" / "github_publication"
DOCX_PATH = OUT_DIR / "Growth_Lab_Core_GitHub_Publication_Guide.docx"
PDF_PATH = OUT_DIR / "Growth_Lab_Core_GitHub_Publication_Guide.pdf"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(8)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "Growth Lab Core GitHub Publication Guide"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "項目"
    table.rows[0].cells[1].text = "内容"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
    set_table_width(table, [2400, 6960])
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_width(table, widths)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def pdf_styles():
    font_path = Path(r"C:\Windows\Fonts\NotoSansJP-VF.ttf")
    if font_path.exists():
        pdfmetrics.registerFont(TTFont("NotoSansJP", str(font_path)))
        font_name = "NotoSansJP"
    else:
        pdfmetrics.registerFont(UnicodeCIDFont("HeiseiKakuGo-W5"))
        font_name = "HeiseiKakuGo-W5"
    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "JPBase",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    return {
        "title": ParagraphStyle(
            "JPTitle",
            parent=base,
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0B2545"),
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "JPSubtitle",
            parent=base,
            fontSize=13,
            leading=17,
            textColor=colors.HexColor("#1F4D78"),
            spaceAfter=8,
        ),
        "meta": ParagraphStyle(
            "JPMeta",
            parent=base,
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#555555"),
            spaceAfter=12,
        ),
        "h1": ParagraphStyle(
            "JPH1",
            parent=base,
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "JPH2",
            parent=base,
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
        ),
        "body": base,
        "small": ParagraphStyle("JPSmall", parent=base, fontSize=8.5, leading=11),
    }


def ptext(text, style):
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def pdf_table(headers, rows, widths, styles):
    data = [[ptext(h, styles["small"]) for h in headers]]
    for row in rows:
        data.append([ptext(c, styles["small"]) for c in row])
    table = Table(data, colWidths=widths, hAlign="CENTER", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C2CC")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_bullets(items, styles):
    return ListFlowable(
        [ListItem(ptext(item, styles["body"]), leftIndent=10) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
    )


def pdf_numbers(items, styles):
    return ListFlowable(
        [ListItem(ptext(item, styles["body"]), leftIndent=10) for item in items],
        bulletType="1",
        leftIndent=18,
    )


def build_pdf():
    styles = pdf_styles()
    story = [
        ptext("Growth Lab Core", styles["title"]),
        ptext("GitHub公開準備ガイド及び仕様書・管理方針", styles["subtitle"]),
        ptext("Version: 1.0 Draft / Created: 2026-07-13 / Owner: Human Owner / Prepared by: Codex", styles["meta"]),
        ptext("1. 文書の目的", styles["h1"]),
        ptext(
            "本資料は、Growth Lab CoreをGitHub上で管理・公開するために必要な情報を、Human Ownerが判断しやすい形で整理した公開準備ガイドである。対象は仕様書、管理方針、公開前チェック、リポジトリ運用、セキュリティ境界、CI品質ゲート、今後の更新運用である。",
            styles["body"],
        ),
        ptext(
            "本資料は公開作業そのものを代行するものではない。GitHubへのログイン、repository作成、push、Recovery codesや認証情報の操作はHuman Ownerが実施し、Codexは仕様整理と文書化を支援する。",
            styles["body"],
        ),
        ptext("2. 公開方針の要約", styles["h1"]),
        pdf_table(
            ["項目", "内容"],
            [
                ("推奨repository名", "growth_lab_core"),
                ("推奨初期公開範囲", "Private repositoryとして開始し、公開可否レビュー完了後に必要範囲を公開する"),
                ("正本フォーマット", "Markdown。WordとPDFはレビュー・共有用の二次成果物"),
                ("公開前の必須条件", "秘密情報、認証情報、個人情報、未承認URL、Recovery codesがrepositoryに含まれないこと"),
                ("CI方針", "OpenAPI Lint workflowは初期段階ではworkflow_dispatchのみ。push/pull_request triggerは後続計画で扱う"),
                ("承認原則", "セキュリティ、公開範囲、CI実行、外部連携はHuman Owner承認後に進める"),
            ],
            [1.75 * inch, 4.65 * inch],
            styles,
        ),
        Spacer(1, 8),
        ptext("3. GitHubに載せる情報", styles["h1"]),
        ptext("公開またはrepository管理の対象にできる情報は、設計・運用・検証に必要で、秘密情報を含まないものに限定する。", styles["body"]),
        pdf_table(
            ["分類", "公開・管理対象", "備考"],
            [
                ("Master Architecture", "architecture/master/00から14のMarkdown", "Growth Lab Coreの上位設計SSOT。Markdownを正本とする。"),
                ("ADR", "architecture/master/adr配下のADR", "重要な設計判断、却下案、見直し条件を記録する。"),
                ("OpenAPI", "implementation/openapi配下のOpenAPI Draft", "dummy URLと安全なplaceholderのみを使用する。"),
                ("Lint/CI設定", ".github/workflows/openapi-lint.yml、Spectral config", "初期は手動実行のみ。権限はcontents: readを基本とする。"),
                ("レビュー記録", "review、test_specifications、test_results", "秘密情報やraw logの無制限転載を避け、summary中心で記録する。"),
                ("進捗管理", "Project_Progress.md、CHANGELOG.md", "作業進捗と変更履歴を追跡する。"),
            ],
            [1.35 * inch, 2.35 * inch, 2.7 * inch],
            styles,
        ),
        Spacer(1, 8),
        ptext("4. GitHubに載せてはいけない情報", styles["h1"]),
        pdf_bullets(
            [
                "Recovery codes、Password、Token、Secret、API Key、Authenticator seedの実体値",
                ".envや本番credential、private registry token、CI secret値",
                "個人を特定できる不要な情報、未承認の実在URL、未承認の外部接続情報",
                "GitHub Actions raw logの全文転載、秘密情報を含む可能性がある添付・スクリーンショット",
                "A8.net、SNS、ASP等の規約についてHuman Ownerまたは専門家確認前の断定的な法務判断",
            ],
            styles,
        ),
        PageBreak(),
        ptext("5. 推奨repository構成", styles["h1"]),
        pdf_table(
            ["Path", "役割", "公開判断"],
            [
                ("architecture/master", "Master Architecture Specification", "公開候補。ただし秘密情報スキャン後。"),
                ("architecture/master/adr", "Architecture Decision Records", "公開候補。判断理由の透明性を高める。"),
                ("implementation/openapi", "Approval Gate等のOpenAPI Draft", "公開候補。placeholderとdummy exampleのみ。"),
                ("implementation/lint_configs", "OpenAPI lint rules", "公開候補。品質ゲートの再現性を担保する。"),
                (".github/workflows", "GitHub Actions workflow", "公開候補。triggerとpermissionの安全性確認後。"),
                ("_backup", "作業前バックアップ", "原則非公開または公開前に整理。古い情報混入に注意。"),
                (".env", "ローカル環境変数", "非公開。Git管理対象外であることを確認。"),
                ("node_modules / .next / .pnpm-store", "生成物・依存キャッシュ", "非公開。repositoryに含めない。"),
            ],
            [1.75 * inch, 2.7 * inch, 1.95 * inch],
            styles,
        ),
        Spacer(1, 8),
        ptext("6. 公開前チェックリスト", styles["h1"]),
        pdf_numbers(
            [
                "Recovery codesがC:\\claudcode_ap\\growth_lab_core外に保管され、git statusに表示されないことを確認する。",
                ".env、credential、token、secret、API key、private registry tokenがGit管理対象外であることを確認する。",
                "公開候補Markdown、OpenAPI、workflow、lint configに文字化け、置換文字、未処理placeholderがないことを確認する。",
                "OpenAPI YAMLが構文検証とSpectral lintを通過していることを確認する。",
                "GitHub Actions workflowが初期段階ではworkflow_dispatchのみであり、contents: read権限のみであることを確認する。",
                "Human Ownerがrepository visibility、repository名、初回push範囲、公開範囲を承認する。",
                "公開後もCHANGELOG、Project_Progress、ADRを継続更新する運用を確認する。",
            ],
            styles,
        ),
        PageBreak(),
        ptext("7. GitHub repository初期設定方針", styles["h1"]),
        pdf_table(
            ["項目", "方針"],
            [
                ("Visibility", "初期はPrivateを推奨。公開レビュー完了後に必要範囲をPublic化する。"),
                ("Repository name", "growth_lab_core"),
                ("初期作成時のREADME等", "既存ローカルrepositoryを正とするため、GitHub側ではREADME、.gitignore、LICENSEを自動生成しない。"),
                ("Default branch", "Human Ownerが決定。mainを候補とする。"),
                ("Branch protection", "初回push後、PR運用とCI品質ゲートの計画に合わせて別途設定する。"),
                ("Secrets / Variables", "必要になるまで作成しない。作成時は用途、権限、rotation、ownerを記録する。"),
            ],
            [2.0 * inch, 4.4 * inch],
            styles,
        ),
        PageBreak(),
        ptext("8. 仕様書管理方針", styles["h1"]),
        pdf_bullets(
            [
                "architecture/master配下のMarkdownをSingle Source of Truthとして扱う。",
                "WordとPDFはレビュー・共有・承認用の出力物であり、正本変更はMarkdownへ反映する。",
                "重要な設計判断はADRとして記録し、判断理由、採用案、却下案、影響範囲、見直し条件を残す。",
                "設計、仕様、実装、設定、依存関係、運用ルールが変わる場合はCHANGELOGを更新する。",
                "単純な進捗更新のみの場合はProject_Progressを更新し、CHANGELOG更新は不要とする。",
            ],
            styles,
        ),
        ptext("9. セキュリティ及び公開境界", styles["h1"]),
        pdf_table(
            ["領域", "公開前の確認", "停止条件"],
            [
                ("Secrets", "実体値が文書・log・設定に含まれない", "1件でも疑いがあれば公開停止"),
                ("OAuth / API", "scope、callback、token保存方針がplaceholderまたは承認済み", "実API credentialが混入した場合"),
                ("CI", "workflow permissionが最小権限", "不明なsecret参照や外部送信がある場合"),
                ("Logs", "summary中心でraw log全文を保存しない", "secret疑いのあるlogを転載した場合"),
                ("Personal data", "不要な個人情報を含めない", "個人を特定できる情報が未承認で含まれる場合"),
            ],
            [1.25 * inch, 2.85 * inch, 2.3 * inch],
            styles,
        ),
        ptext("10. CI品質ゲート方針", styles["h1"]),
        ptext(
            "OpenAPI lintはGitHub Actionsで管理する。ただし初期段階ではHuman Ownerが手動でworkflow_dispatchを実行し、結果を共有した後にCodexがレポート化する。CodexはGitHub Actionsを直接実行せず、GitHub CLIやGitHub Web UIを操作しない。",
            styles["body"],
        ),
        pdf_bullets(
            [
                "workflow name: OpenAPI Lint",
                "job name: OpenAPI Spectral Lint",
                "initial trigger: workflow_dispatch only",
                "permission: contents: read only",
                "quality command: pnpm run lint:openapi",
                "version command: pnpm run lint:openapi:version",
            ],
            styles,
        ),
        PageBreak(),
        ptext("11. 役割分担", styles["h1"]),
        pdf_table(
            ["Role", "責任範囲", "主な判断"],
            [
                ("Human Owner", "GitHub作成、公開範囲、認証、secret、外部連携の最終判断", "Public化、初回push、CI実行、Recovery codes保管"),
                ("Codex", "仕様整理、文書化、ローカル検証、レポート作成", "外部サービス操作は行わず、共有結果を文書化"),
                ("Architecture Documents", "SSOTとして設計と運用方針を保持", "変更時はADRとCHANGELOGで追跡"),
                ("CI / Lint", "OpenAPI品質確認と将来のPR品質ゲート", "初期は手動、後続でPR triggerを計画"),
            ],
            [1.45 * inch, 3.25 * inch, 1.7 * inch],
            styles,
        ),
        ptext("12. Human Owner向け最終確認", styles["h1"]),
        pdf_numbers(
            [
                "Recovery codesをrepository外の安全な場所へ移動済みである。",
                "Recovery codesのファイル名、内容、保存先詳細をrepository内文書へ記録していない。",
                "git statusに秘密情報らしきファイルが表示されていない。",
                "GitHub repositoryはPrivateで作成する。",
                "GitHub側でREADME、.gitignore、LICENSEを自動生成しない。",
                "初回push前にsecret scanと公開対象レビューを行う。",
                "OpenAPI Lint workflowの手動実行結果をCodexへ共有し、実行レポートを作成する。",
            ],
            styles,
        ),
        ptext("13. 次の推奨作業", styles["h1"]),
        pdf_bullets(
            [
                "Recovery codesのrepository外保管確認",
                "git statusによる秘密情報混入確認",
                "Private repository作成",
                "初回push対象の最終レビュー",
                "OpenAPI Lint workflowのGitHub Actions表示確認",
                "Human Ownerによるmanual workflow execution",
                "GitHub Actions Manual Workflow Execution Report作成",
                "Pull request trigger planning",
            ],
            styles,
        ),
        PageBreak(),
        ptext("Appendix A. 公開可否の簡易判定表", styles["h1"]),
        pdf_table(
            ["確認項目", "公開可", "公開停止"],
            [
                ("設計書", "秘密情報がなく、Markdown正本と一致", "未レビューのcredential、未承認URL、文字化けあり"),
                ("ADR", "判断理由と影響範囲のみ", "個人情報、secret、外部契約の断定的法務判断あり"),
                ("OpenAPI", "dummy URLとplaceholderのみ", "実API endpoint、実token、実secretあり"),
                ("CI", "最小権限、manual trigger", "不明なsecret参照、過大権限、未承認triggerあり"),
                ("Logs", "要約と件数のみ", "raw log全文、secret疑いの値あり"),
            ],
            [2.0 * inch, 2.2 * inch, 2.2 * inch],
            styles,
        ),
    ]
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="Growth Lab Core GitHub Publication Guide",
    )
    doc.build(story)
    print(PDF_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)

    doc.add_paragraph("Growth Lab Core", style="Title")
    subtitle = doc.add_paragraph("GitHub公開準備ガイド及び仕様書・管理方針")
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor(31, 77, 120)
    meta = doc.add_paragraph("Version: 1.0 Draft / Created: 2026-07-13 / Owner: Human Owner / Prepared by: Codex")
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_heading("1. 文書の目的", level=1)
    doc.add_paragraph(
        "本資料は、Growth Lab CoreをGitHub上で管理・公開するために必要な情報を、"
        "Human Ownerが判断しやすい形で整理した公開準備ガイドである。"
        "対象は仕様書、管理方針、公開前チェック、リポジトリ運用、セキュリティ境界、"
        "CI品質ゲート、今後の更新運用である。"
    )
    doc.add_paragraph(
        "本資料は公開作業そのものを代行するものではない。GitHubへのログイン、repository作成、"
        "push、Recovery codesや認証情報の操作はHuman Ownerが実施し、Codexは仕様整理と文書化を支援する。"
    )

    doc.add_heading("2. 公開方針の要約", level=1)
    add_kv_table(
        doc,
        [
            ("推奨repository名", "growth_lab_core"),
            ("推奨初期公開範囲", "Private repositoryとして開始し、公開可否レビュー完了後に必要範囲を公開する"),
            ("正本フォーマット", "Markdown。WordとPDFはレビュー・共有用の二次成果物"),
            ("公開前の必須条件", "秘密情報、認証情報、個人情報、未承認URL、Recovery codesがrepositoryに含まれないこと"),
            ("CI方針", "OpenAPI Lint workflowは初期段階ではworkflow_dispatchのみ。push/pull_request triggerは後続計画で扱う"),
            ("承認原則", "セキュリティ、公開範囲、CI実行、外部連携はHuman Owner承認後に進める"),
        ],
    )

    doc.add_heading("3. GitHubに載せる情報", level=1)
    doc.add_paragraph("公開またはrepository管理の対象にできる情報は、設計・運用・検証に必要で、秘密情報を含まないものに限定する。")
    add_table(
        doc,
        ["分類", "公開・管理対象", "備考"],
        [
            ("Master Architecture", "architecture/master/00から14のMarkdown", "Growth Lab Coreの上位設計SSOT。Markdownを正本とする。"),
            ("ADR", "architecture/master/adr配下のADR", "重要な設計判断、却下案、見直し条件を記録する。"),
            ("OpenAPI", "implementation/openapi配下のOpenAPI Draft", "dummy URLと安全なplaceholderのみを使用する。"),
            ("Lint/CI設定", ".github/workflows/openapi-lint.yml、Spectral config", "初期は手動実行のみ。権限はcontents: readを基本とする。"),
            ("レビュー記録", "review、test_specifications、test_results", "秘密情報やraw logの無制限転載を避け、summary中心で記録する。"),
            ("進捗管理", "Project_Progress.md、CHANGELOG.md", "作業進捗と変更履歴を追跡する。"),
        ],
        [2100, 3460, 3800],
    )

    doc.add_heading("4. GitHubに載せてはいけない情報", level=1)
    add_bullets(
        doc,
        [
            "Recovery codes、Password、Token、Secret、API Key、Authenticator seedの実体値",
            ".envや本番credential、private registry token、CI secret値",
            "個人を特定できる不要な情報、未承認の実在URL、未承認の外部接続情報",
            "GitHub Actions raw logの全文転載、秘密情報を含む可能性がある添付・スクリーンショット",
            "A8.net、SNS、ASP等の規約についてHuman Ownerまたは専門家確認前の断定的な法務判断",
        ],
    )
    doc.add_paragraph(
        "認証情報らしきファイルを検出した場合は、内容を開かず、必要最小限の情報のみHuman Owner確認事項として扱う。"
    )

    doc.add_heading("5. 推奨repository構成", level=1)
    add_table(
        doc,
        ["Path", "役割", "公開判断"],
        [
            ("architecture/master", "Master Architecture Specification", "公開候補。ただし秘密情報スキャン後。"),
            ("architecture/master/adr", "Architecture Decision Records", "公開候補。判断理由の透明性を高める。"),
            ("implementation/openapi", "Approval Gate等のOpenAPI Draft", "公開候補。placeholderとdummy exampleのみ。"),
            ("implementation/lint_configs", "OpenAPI lint rules", "公開候補。品質ゲートの再現性を担保する。"),
            (".github/workflows", "GitHub Actions workflow", "公開候補。triggerとpermissionの安全性確認後。"),
            ("_backup", "作業前バックアップ", "原則非公開または公開前に整理。古い情報混入に注意。"),
            (".env", "ローカル環境変数", "非公開。Git管理対象外であることを確認。"),
            ("node_modules / .next / .pnpm-store", "生成物・依存キャッシュ", "非公開。repositoryに含めない。"),
        ],
        [2600, 3960, 2800],
    )

    doc.add_heading("6. 公開前チェックリスト", level=1)
    add_numbers(
        doc,
        [
            "Recovery codesがC:\\claudcode_ap\\growth_lab_core外に保管され、git statusに表示されないことを確認する。",
            ".env、credential、token、secret、API key、private registry tokenがGit管理対象外であることを確認する。",
            "公開候補Markdown、OpenAPI、workflow、lint configに文字化け、置換文字、未処理placeholderがないことを確認する。",
            "OpenAPI YAMLが構文検証とSpectral lintを通過していることを確認する。",
            "GitHub Actions workflowが初期段階ではworkflow_dispatchのみであり、contents: read権限のみであることを確認する。",
            "Human Ownerがrepository visibility、repository名、初回push範囲、公開範囲を承認する。",
            "公開後もCHANGELOG、Project_Progress、ADRを継続更新する運用を確認する。",
        ],
    )

    doc.add_heading("7. GitHub repository初期設定方針", level=1)
    add_kv_table(
        doc,
        [
            ("Visibility", "初期はPrivateを推奨。公開レビュー完了後に必要範囲をPublic化する。"),
            ("Repository name", "growth_lab_core"),
            ("初期作成時のREADME等", "既存ローカルrepositoryを正とするため、GitHub側ではREADME、.gitignore、LICENSEを自動生成しない。"),
            ("Default branch", "Human Ownerが決定。mainを候補とする。"),
            ("Branch protection", "初回push後、PR運用とCI品質ゲートの計画に合わせて別途設定する。"),
            ("Secrets / Variables", "必要になるまで作成しない。作成時は用途、権限、rotation、ownerを記録する。"),
        ],
    )

    doc.add_heading("8. 仕様書管理方針", level=1)
    add_bullets(
        doc,
        [
            "architecture/master配下のMarkdownをSingle Source of Truthとして扱う。",
            "WordとPDFはレビュー・共有・承認用の出力物であり、正本変更はMarkdownへ反映する。",
            "重要な設計判断はADRとして記録し、判断理由、採用案、却下案、影響範囲、見直し条件を残す。",
            "設計、仕様、実装、設定、依存関係、運用ルールが変わる場合はCHANGELOGを更新する。",
            "単純な進捗更新のみの場合はProject_Progressを更新し、CHANGELOG更新は不要とする。",
        ],
    )

    doc.add_heading("9. セキュリティ及び公開境界", level=1)
    add_table(
        doc,
        ["領域", "公開前の確認", "停止条件"],
        [
            ("Secrets", "実体値が文書・log・設定に含まれない", "1件でも疑いがあれば公開停止"),
            ("OAuth / API", "scope、callback、token保存方針がplaceholderまたは承認済み", "実API credentialが混入した場合"),
            ("CI", "workflow permissionが最小権限", "不明なsecret参照や外部送信がある場合"),
            ("Logs", "summary中心でraw log全文を保存しない", "secret疑いのあるlogを転載した場合"),
            ("Personal data", "不要な個人情報を含めない", "個人を特定できる情報が未承認で含まれる場合"),
        ],
        [1800, 4160, 3400],
    )

    doc.add_heading("10. CI品質ゲート方針", level=1)
    doc.add_paragraph(
        "OpenAPI lintはGitHub Actionsで管理する。ただし初期段階ではHuman Ownerが手動でworkflow_dispatchを実行し、"
        "結果を共有した後にCodexがレポート化する。CodexはGitHub Actionsを直接実行せず、GitHub CLIやGitHub Web UIを操作しない。"
    )
    add_bullets(
        doc,
        [
            "workflow name: OpenAPI Lint",
            "job name: OpenAPI Spectral Lint",
            "initial trigger: workflow_dispatch only",
            "permission: contents: read only",
            "quality command: pnpm run lint:openapi",
            "version command: pnpm run lint:openapi:version",
        ],
    )

    doc.add_heading("11. 役割分担", level=1)
    add_table(
        doc,
        ["Role", "責任範囲", "主な判断"],
        [
            ("Human Owner", "GitHub作成、公開範囲、認証、secret、外部連携の最終判断", "Public化、初回push、CI実行、Recovery codes保管"),
            ("Codex", "仕様整理、文書化、ローカル検証、レポート作成", "外部サービス操作は行わず、共有結果を文書化"),
            ("Architecture Documents", "SSOTとして設計と運用方針を保持", "変更時はADRとCHANGELOGで追跡"),
            ("CI / Lint", "OpenAPI品質確認と将来のPR品質ゲート", "初期は手動、後続でPR triggerを計画"),
        ],
        [2100, 4760, 2500],
    )

    doc.add_heading("12. Human Owner向け最終確認", level=1)
    add_numbers(
        doc,
        [
            "Recovery codesをrepository外の安全な場所へ移動済みである。",
            "Recovery codesのファイル名、内容、保存先詳細をrepository内文書へ記録していない。",
            "git statusに秘密情報らしきファイルが表示されていない。",
            "GitHub repositoryはPrivateで作成する。",
            "GitHub側でREADME、.gitignore、LICENSEを自動生成しない。",
            "初回push前にsecret scanと公開対象レビューを行う。",
            "OpenAPI Lint workflowの手動実行結果をCodexへ共有し、実行レポートを作成する。",
        ],
    )

    doc.add_heading("13. 次の推奨作業", level=1)
    add_bullets(
        doc,
        [
            "Recovery codesのrepository外保管確認",
            "git statusによる秘密情報混入確認",
            "Private repository作成",
            "初回push対象の最終レビュー",
            "OpenAPI Lint workflowのGitHub Actions表示確認",
            "Human Ownerによるmanual workflow execution",
            "GitHub Actions Manual Workflow Execution Report作成",
            "Pull request trigger planning",
        ],
    )

    doc.add_page_break()
    doc.add_heading("Appendix A. 公開可否の簡易判定表", level=1)
    add_table(
        doc,
        ["確認項目", "公開可", "公開停止"],
        [
            ("設計書", "秘密情報がなく、Markdown正本と一致", "未レビューのcredential、未承認URL、文字化けあり"),
            ("ADR", "判断理由と影響範囲のみ", "個人情報、secret、外部契約の断定的法務判断あり"),
            ("OpenAPI", "dummy URLとplaceholderのみ", "実API endpoint、実token、実secretあり"),
            ("CI", "最小権限、manual trigger", "不明なsecret参照、過大権限、未承認triggerあり"),
            ("Logs", "要約と件数のみ", "raw log全文、secret疑いの値あり"),
        ],
        [3000, 3180, 3180],
    )

    doc.save(DOCX_PATH)
    build_pdf()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
